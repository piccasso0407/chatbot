import streamlit as st
import os

st.set_page_config(layout="wide")

current_dir = os.path.dirname(os.path.abspath(__file__))
css_path = os.path.join(current_dir, 'style.css')

# CSS 파일 로드 함수
def load_css(file_name):
    if os.path.exists(file_name):
        with open(file_name, 'r', encoding='utf-8') as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    else:
        st.error(f"CSS 파일을 찾을 수 없습니다: {file_name}")

# CSS 파일 로드
load_css(css_path)

# Pretendard 폰트 로드
st.markdown("""
    <link href="https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css" rel="stylesheet">
    <style>
        html, body, [class*="css"] {
            font-family: 'Pretendard', sans-serif;
        }
    </style>
    """, unsafe_allow_html=True)

# 익스펜더 안에 코드 표시
with st.expander("코드 보기"):
    code = '''
    # Streamlit 앱의 제목
    st.title("저염식 식단 챗봇")

    # 필요한 라이브러리 임포트
    import time
    import pandas as pd
    import requests
    from PyPDF2 import PdfReader
    import docx
    from langchain.document_loaders import PyPDFLoader, Docx2txtLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.embeddings import HuggingFaceEmbeddings
    from langchain.vectorstores import FAISS
    from langchain.memory import ConversationBufferMemory
    from langchain.schema import Document

    # PDF 파일 읽기 함수
    def read_pdf(file_path):
        pdf_text = []
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                pdf_text.append(page.extract_text())
        return pdf_text

    # DOCX 파일 읽기 함수
    def read_docx(file_path):
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\\n'.join(full_text)

    # 파일 경로
    docx_file_path = "./data/recipes.docx"
    pdf_file_path2 = "./1.pdf"

    # PDF 및 DOCX 파일에서 텍스트 불러오기
    docx_content = read_docx(docx_file_path)
    pdf_content2 = read_pdf(pdf_file_path2)

    # 데이터를 딕셔너리 형태로 정리
    data = {
        "항목": [
            "이름", "만나이", "성별", "키", "체중", "허리둘레",
            "이상지질혈증 여부", "당뇨병 여부",
            "음주 여부", "음주 빈도", "1회 주량",
            "고강도 운동 여부", "중강도 운동 여부",
            "걷기, 자전거 운동", "1주일", "총 운동 시간",
            "음주 점수", "신체활동 점수", "고혈압 확률",
        ],
        "값": [
            "김첨지", 38, "남자", "173.0 cm", "89.0 kg", "88.9 cm",
            "없음", "없음",
            "마신다", "월1회정도", "1-2잔",
            "안 한다", "안 한다",
            "한다", "5일", "2시간 0분",
            "9점", "6.39859점", "21.41%",
        ]
    }

    # 데이터프레임 생성
    df = pd.DataFrame(data)
    df1 = df.reset_index(drop=True)

    # 사이드바에 데이터프레임 표시
    st.sidebar.title("김첨지님의 건강 보고서")
    st.sidebar.dataframe(df1)

    # clear_all 함수 정의
    def clear_all():
        # 모든 캐시 지우기
        st.cache_data.clear()
        st.cache_resource.clear()

        # 세션 상태 초기화
        for key in list(st.session_state.keys()):
            del st.session_state[key]

        st.success("모든 캐시와 저장된 상태가 초기화되었습니다!")
        st.experimental_rerun()  # 페이지 새로고침

    # 초기화 버튼을 사이드바로 이동
    custom_css = """
    <style>
        .stButton>button {
            background-color: rgb(255, 241, 219);
            color: black;
            border: 1px solid rgba(0,0,0,0.1);
        }
        .stButton>button:hover {
            background-color: rgba(255, 241, 219, 0.8);
            color: black;
        }
    </style>
    """

    # 커스텀 CSS 적용
    st.sidebar.markdown(custom_css, unsafe_allow_html=True)

    # 초기화 버튼을 사이드바로 이동
    if st.sidebar.button('모든 데이터 초기화'):
        clear_all()

    documents = []

    # docx_content는 하나의 문자열이므로 바로 추가
    if isinstance(docx_content, str):
        documents.append(Document(page_content=docx_content))
    else:
        st.error("DOCX 내용이 올바른 문자열 형식이 아닙니다.")

    # pdf_content2도 리스트일 가능성이 있으므로 개별 텍스트로 변환
    for text in pdf_content2:
        if isinstance(text, str):
            documents.append(Document(page_content=text))
        else:
            st.error("두 번째 PDF 내용이 올바른 문자열 형식이 아닙니다.")

    # 텍스트를 청크로 분리하는 함수
    def get_text_chunks(documents):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=700,
            chunk_overlap=100,
            length_function=len
        )
        chunks = text_splitter.split_documents(documents)
        return chunks

    # 벡터 스토어 생성 함수
    def get_vectorstore(_text_chunks):
        embeddings = HuggingFaceEmbeddings(
            model_name="jhgan/ko-sroberta-multitask",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        vectordb = FAISS.from_documents(_text_chunks, embeddings)
        return vectordb

    # 시스템 프롬프트 템플릿 생성
    def get_system_message(user_info):
        system_message = f"""
    당신은 건강 전문가입니다. 사용자의 건강 정보를 바탕으로 맞춤형 식단과 건강 조언을 제공하는 역할을 합니다.
    제공된 정보를 활용하여 사용자의 질문에 정확하고 유용한 답변을 해주세요.

    사용자 정보:
    이름: {user_info['이름']}
    나이: {user_info['만나이']}세
    성별: {user_info['성별']}
    키: {user_info['키']}
    체중: {user_info['체중']}
    허리둘레: {user_info['허리둘레']}
    이상지질혈증: {user_info['이상지질혈증 여부']}
    당뇨병: {user_info['당뇨병 여부']}
    음주 여부: {user_info['음주 여부']}
    음주 빈도: {user_info['음주 빈도']}
    1회 주량: {user_info['1회 주량']}
    고강도 운동: {user_info['고강도 운동 여부']}
    중강도 운동: {user_info['중강도 운동 여부']}
    걷기/자전거: {user_info['걷기, 자전거 운동']} (주 {user_info['1주일']}일, {user_info['총 운동 시간']})
    음주 점수: {user_info['음주 점수']}
    신체활동 점수: {user_info['신체활동 점수']}
    고혈압 확률: {user_info['고혈압 확률']}

    이 정보를 바탕으로 사용자에게 맞춤형 건강 조언과 식단을 제공해주세요.
    """
        return system_message

    user_info = {
        "이름": "김첨지",
        "만나이": 38,
        "성별": "남자",
        "키": "173.0 cm",
        "체중": "89.0 kg",
        "허리둘레": "88.9 cm",
        "이상지질혈증 여부": "없음",
        "당뇨병 여부": "없음",
        "음주 여부": "마신다",
        "음주 빈도": "월1회정도",
        "1회 주량": "1-2잔",
        "고강도 운동 여부": "안 한다",
        "중강도 운동 여부": "안 한다",
        "걷기, 자전거 운동": "한다",
        "1주일": "5일",
        "총 운동 시간": "2시간 0분",
        "음주 점수": "9점",
        "신체활동 점수": "6.39859점",
        "고혈압 확률": "21.41%"
    }

    system_message = get_system_message(user_info)

    # 텍스트 청크 생성
    text_chunks = get_text_chunks(documents)

    # 벡터 스토어 생성
    vectorstore = get_vectorstore(text_chunks)
    LMSTUDIO_URL = "http://localhost:1234/v1/chat/completions"

    # 메모리 초기화
    memory = ConversationBufferMemory(return_messages=True)

    # 요청 보내기
    def chat_with_bot(question, vector_store):
        retriever = vector_store.as_retriever(search_kwargs={"k": 7})
        docs = retriever.get_relevant_documents(question)

        context = "\\n\\n".join([doc.page_content for doc in docs])

        prompt = get_prompt_template(question, context)

        headers = {"Content-Type": "application/json"}
        data = {
            "messages": [
                {"role": "system", "content": system_message},
                {"role": "user", "content": prompt}
            ],
            "model": "teddylee777/EEVE-Korean-Instruct-10.8B-v1.0-gguf/EEVE-Korean-Instruct-10.8B-v1.0-Q5_K_M.gguf",
            "temperature": 0.3
        }
        response = requests.post(LMSTUDIO_URL, headers=headers, json=data)

        if response.status_code == 200:
            bot_response = response.json()['choices'][0]['message']['content']

            memory.chat_memory.add_user_message(question)
            memory.chat_memory.add_ai_message(bot_response)

            return bot_response, docs
        else:
            return f"오류 발생: {response.status_code}", []

    # 대화형 체인 생성
    def get_conversation_chain(vectorstore):
        return lambda x: chat_with_bot(x["question"], vectorstore)

    # 대화형 체인 생성
    conversation_chain = get_conversation_chain(vectorstore)

    if 'conversation' not in st.session_state:
        st.session_state.conversation = conversation_chain

    def get_prompt_template(query, data):
        base_prompt = f"""
    참고 정보:
    {data}

    질문: {query}

    지침:
    - 제공된 정보를 바탕으로 질문에 정확하고 구체적으로 답변해 주세요.
    - 모든 링크는 제거해 주세요.
    - 나트륨 함량이 500mg 이하인 음식만 추천.
    """

        if "식단" in query:
            return base_prompt + """
    응답 형식:
    1. 추천 식단 (아침, 점심, 저녁)
    2. 각 식단의 영양 정보
    3. 식단 선택 이유 간단히 설명
    """
        elif "레시피" in query:
            return base_prompt + """
    응답 형식:
    1. 재료 목록
    2. 조리 단계
    3. 영양 정보
    """
        else:
            return base_prompt + """
    응답 형식:
    1. 질문에 대한 직접적인 답변
    2. 추가 설명 또는 관련 정보 (필요시)
    3. 주의사항 또는 권고사항 (적절한 경우)
    """

    # 채팅 로직
    if 'messages' not in st.session_state:
        st.session_state['messages'] = []

    query = st.chat_input("식단 또는 레시피를 물어보세요.")
    if query:
        st.session_state['messages'].append({"role": "user", "content": query})
        with st.spinner("Thinking..."):
            try:
                chain = st.session_state.conversation
                response, source_documents = chain({"question": query})

                if not response:
                    st.error("응답을 받지 못했습니다. 다시 시도해 주세요.")
                else:
                    st.session_state['messages'].append({"role": "assistant", "content": response})

                    # 채팅 메시지 표시
                    for message in st.session_state['messages']:
                        if message['role'] == 'user':
                            st.chat_message("user").markdown(message['content'])
                        else:
                            st.chat_message("assistant").markdown(message['content'])

                    if source_documents:
                        with st.expander("참고 문서 확인"):
                            for doc in source_documents:
                                st.markdown(doc.metadata.get('source', ''), help=doc.page_content)
            except Exception as e:
                st.error(f"오류가 발생했습니다: {str(e)}")
                st.error(f"프롬프트: {query}")
    '''
    st.code(code, language='python')
current_dir = os.path.dirname(os.path.abspath(__file__))
image_path2 = os.path.join(current_dir, "images", "thinking.jpg")
image_path3 = os.path.join(current_dir, "images", "answer.jpg")
image_path4 = os.path.join(current_dir, "images", "chamgo.jpg")

st.image(image_path2)
st.image(image_path3)
st.image(image_path4)


# 비디오 파일의 절대경로를 지정합니다.
video_file = os.path.join(current_dir, "images", "bandicam 2024-09-30 15-39-46-483.mp4")  # 실제 경로로 변경하세요.

# 비디오 파일이 존재하는지 확인합니다.
if os.path.exists(video_file):
    with open(video_file, 'rb') as f:
        video_bytes = f.read()
    # 비디오를 표시합니다.
    st.video(video_bytes)
else:
    st.error('비디오 파일을 찾을 수 없습니다.')
