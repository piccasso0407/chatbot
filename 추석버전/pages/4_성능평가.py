from langchain_community.document_loaders import WebBaseLoader
from langchain_community.vectorstores import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.chat_models import ChatOllama
from langchain.prompts import ChatPromptTemplate
from pprint import pprint
import streamlit as st
from PyPDF2 import PdfReader
import docx
import streamlit as st
import tiktoken
from loguru import logger
import time
import concurrent.futures
import pandas as pd
import json
import os

from langchain.chains import ConversationalRetrievalChain
from langchain.document_loaders import PyPDFLoader, Docx2txtLoader, UnstructuredPowerPointLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.vectorstores import FAISS
from langchain.memory import StreamlitChatMessageHistory
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain_openai import ChatOpenAI

import tiktoken
from loguru import logger
import time
import concurrent.futures
import pandas as pd
import json
import os
from transformers import AutoTokenizer
from langchain.chains import ConversationalRetrievalChain
from langchain.document_loaders import PyPDFLoader, Docx2txtLoader, UnstructuredPowerPointLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.vectorstores import FAISS
from langchain.memory import StreamlitChatMessageHistory
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain_openai import ChatOpenAI
st.set_page_config(layout="wide")

current_dir = os.path.dirname(os.path.abspath(__file__))
css_path = os.path.join(current_dir, 'style.css')

# CSS 파일 로드 함수
def load_css(file_name):
    if os.path.exists(file_name):
        with open(file_name, 'r', encoding='utf-8') as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    else:
        st.error(f"CSS 파일을 찾을 수 없습니다: {file_name}")

# CSS 파일 로드
load_css(css_path)

# Pretendard 폰트 로드
st.markdown("""
    <link href="https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css" rel="stylesheet">
    <style>
        html, body, [class*="css"] {
            font-family: 'Pretendard', sans-serif;
        }
    </style>
    """, unsafe_allow_html=True)


st.markdown("* * *")
st.subheader("|QnA set")

# Set the current directory and file path
current_dir = os.path.dirname(os.path.abspath(__file__))
file_path = os.path.join(current_dir, 'qnaset1.xlsx')

# Check if the file exists
if os.path.exists(file_path):
    # Read the Excel file into a DataFrame
    qa_data = pd.read_excel(file_path)
    # Display the DataFrame
    st.dataframe(qa_data)

import PyPDF2
import docx  # For reading DOCX files
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain_openai import ChatOpenAI
from langchain.schema import Document
from string import Template
import re
import os
import time
from langchain.prompts import ChatPromptTemplate
import pandas as pd
import streamlit as st

# # 외부 CSS 불러오기
# load_css('./style.css')
# # 폰트 설정
# load_local_font('Pretendard', 'C:\Windows\Fonts\Arial.ttf')

st.markdown("* * *")    
st.subheader("|내 모델 로드")     
with st.expander("코드 보기"):
    st.code('''
# PDF 파일 읽기 함수
def read_pdf(file_path):
    pdf_text = []
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            pdf_text.append(page.extract_text())
    return pdf_text

# DOCX 파일 읽기 함수
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# 파일 경로
pdf_file_path = "./02-우리_몸이_원하는_삼삼한밥상_Ⅸ_본문(화면용_펼침).pdf"
docx_file_path = "./data/recipes.docx"
pdf_file_path2 = "./1.pdf"

# PDF 및 DOCX 파일에서 텍스트 불러오기
pdf_content = read_pdf(pdf_file_path)
docx_content = read_docx(docx_file_path)
pdf_content2 = read_pdf(pdf_file_path2)

# 데이터를 딕셔너리 형태로 정리
data = {
    "항목": [
        "이름", "만나이", "성별", "키", "체중", "허리둘레", 
        "이상지질혈증 여부", "당뇨병 여부", 
        "음주 여부", "음주 빈도", "1회 주량", 
        "폭음 횟수", "절주 권고", 
        "음주 상담 유무",
        "고강도 운동 여부", "1주일", "1회 운동 시간", 
        "중강도 운동 여부", "1주일", "1회 운동 시간?", 
        "걷기, 자전거 운동", "1주일", "총 운동 시간",
        "음주 점수", "신체활동 점수", "고혈압 확률",
    ],
    "값": [
        "김첨지", 38, "남자", "173.0 cm", "89.0 kg", "88.9 cm",
        "없음", "없음",
        "마신다", "월1회정도", "1-2잔",
        "월 1회 미만", "없음", "없음",
        "안 한다", "None", "None", 
        "안 한다", "None", "None",
        "한다", "5일", "2시간 0분",
        "9점", "6.39859점", "21.41%",
    ]
}

# 데이터프레임 생성
df = pd.DataFrame(data)

# PDF 및 DOCX 데이터를 문서 리스트로 변환
documents = []

# pdf_content는 리스트 형태일 가능성이 높으므로 개별 텍스트로 변환
for text in pdf_content:
    if isinstance(text, str):
        documents.append(Document(page_content=text))

# docx_content는 하나의 문자열이므로 바로 추가
if isinstance(docx_content, str):
    documents.append(Document(page_content=docx_content))

# pdf_content2도 리스트일 가능성이 있으므로 개별 텍스트로 변환
for text in pdf_content2:
    if isinstance(text, str):
        documents.append(Document(page_content=text))

# 텍스트를 청크로 분리하는 함수
def get_text_chunks(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size= 100,  # 청크 크기를 더 줄임
        chunk_overlap=5,
        length_function=len
    )
    chunks = text_splitter.split_documents(documents)
    return chunks

# 벡터 스토어 생성 함수
def get_vectorstore(_text_chunks):
    embeddings = HuggingFaceEmbeddings(
        model_name="jhgan/ko-sroberta-multitask",
        model_kwargs={'device': 'cuda'},
        encode_kwargs={'normalize_embeddings': True}
    )
    vectordb = FAISS.from_documents(_text_chunks, embeddings)
    return vectordb

# 시스템 프롬프트 템플릿 생성
system_message = """
data = {
    "항목": [
        "이름", "만나이", "성별", "키", "체중", "허리둘레", 
        "이상지질혈증 여부", "당뇨병 여부", 
        "술을 마십니까?", "술을 얼마나 자주 마십니까?", "한 번에 술을 얼마나 마십니까?", 
        "한 번의 술자리에서 7잔 이상을 마시는 횟수", "술을 끊거나 줄이라는 권고를 받은 적이 있습니까?", 
        "최근 1년 동안 음주 문제로 상담을 받아본 적이 있습니까?",
        "고강도 운동 여부", "1주일에 며칠 하십니까?", "한 번 할 때 몇 시간 하십니까?", 
        "중강도 운동 여부", "1주일에 며칠 하십니까?", "한 번 할 때 몇 시간 하십니까?", 
        "걷기나 자전거를 이용하십니까?", "1주일에 며칠 하십니까?", "대략 몇 시간 움직이십니까?",
        "음주 점수", "신체활동 점수", "고혈압 확률"
    ],
    "값": [
        "김첨지", 38, "남자", "173.0 cm", "89.0 kg", "88.9 cm",
        "없음", "없음",
        "마신다", "월1회정도", "1-2잔",
        "월 1회 미만", "없음", "없음",
        "안 한다", "None", "None", 
        "안 한다", "None", "None",
        "한다", "5일", "2시간 0분",
        "9점", "6.39859점", "21.41%",
    ]
}
너는 이 데이터를 읽고 데이터에 알맞은 식단이나 레시피를 소개하주는 영양전문가야.
"""

# 텍스트 청크 생성
text_chunks = get_text_chunks(documents)

# 벡터 스토어 생성
vectorstore = get_vectorstore(text_chunks)

# 대화형 체인 생성 함수
def get_conversation_chain(_vectorstore):
    llm = ChatOpenAI(
        base_url="http://localhost:1234/v1",
        api_key="lm-studio",
        model="teddylee777/EEVE-Korean-Instruct-10.8B-v1.0-gguf",
        temperature=0.3,
        streaming=True
    )

    retriever = _vectorstore.as_retriever(
        search_type='similarity',
        search_kwargs={"k": 5},
        verbose=False
    )

    # 시스템 메시지를 포함하는 프롬프트 템플릿 생성
    prompt_template = ChatPromptTemplate.from_messages([
        ("system", system_message),
        ("user", "{question}")
    ])

    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm, 
        chain_type="stuff",
        retriever=retriever, 
        return_source_documents=True,
        verbose=False
    )
    return conversation_chain

# 대화형 체인 생성
conversation_chain = get_conversation_chain(vectorstore)
''')
st.subheader("|RAGAS에 필요한 내 llm 답변 출력")       
st.code('''
# 사용자 질의 처리 함수
def process_user_query(query):
    formatted_query = get_prompt_template(query)
    result = conversation_chain({"question": formatted_query, "chat_history": []})
    response = result['answer']
    return response


# 대화 체인을 통해 질의에 대한 답변과 관련 문서를 한 번에 가져오는 함수
def process_question(question):
    # conversation_chain.invoke를 한 번만 호출
    result = conversation_chain.invoke({
        "question": question,
        "chat_history": []
    })
    
    # answer와 contexts를 추출
    answer = result['answer']  # 오직 'answer'만 추출
    
    # 질문이 답변에 포함되어 있으면 제거
    if question in answer:
        answer = answer.replace(question, '').strip()  # 질문 제거 및 양쪽 공백 제거
    
    contexts = [doc.page_content for doc in result['source_documents']]
    
    return answer, contexts

# qa_data의 첫 5개의 질문에 대해서만 'answer'와 'contexts' 열을 채움
qa_data[['answer', 'contexts']] = qa_data['question'].apply(lambda x: pd.Series(process_question(x)))

# Pandas DataFrame을 Dataset으로 변환
from datasets import Dataset

dataset = Dataset.from_pandas(qa_data)
''')

test_middle = pd.read_excel('test_middle.xlsx')
st.dataframe(test_middle)

st.markdown("* * *")
st.subheader("RAGAS: Automated Evaluation of Retrieval Augmented Generation")
st.code('''
from ragas.metrics import (
    faithfulness,
    answer_relevancy,
    context_recall,
    context_precision,
)

from ragas import evaluate
from ragas.metrics import faithfulness, answer_relevancy, context_recall, context_precision
from langchain.llms import OpenAI
from ragas import evaluate
from langchain.embeddings.openai import OpenAIEmbeddings

# OpenAI 모델 사용
langchain_llm = OpenAI(model="gpt-3.5-turbo-instruct", openai_api_key="sk-proj-")
langchain_embeddings = OpenAIEmbeddings(openai_api_key="sk-proj-")
# 메트릭을 하나씩 실행해보는 코드

result_faithfulness = evaluate(
    dataset,
    metrics=[faithfulness],
    llm=langchain_llm,
    embeddings=langchain_embeddings,
    raise_exceptions=False,
)

result_answer_relevancy = evaluate(
    dataset,
    metrics=[answer_relevancy],
    llm=langchain_llm,
    embeddings=langchain_embeddings,
    raise_exceptions=False,
)


result_context_recall = evaluate(
    dataset,
    metrics=[context_recall],
    llm=langchain_llm,
    embeddings=langchain_embeddings,
    raise_exceptions=False,
)


result_context_precision = evaluate(
    dataset,
    metrics=[context_precision],
    llm=langchain_llm,
    embeddings=langchain_embeddings,
    raise_exceptions=False,
)


# 결과 출력
print("Faithfulness Result:", result_faithfulness)
print("Answer Relevancy Result:", result_answer_relevancy)
print("Context Recall Result:", result_context_recall)
print("Context Precision Result:", result_context_precision)
        
''')
st.markdown("* * *")
st.write("참고문헌")
# Arxiv 요약 페이지 링크 추가
st.markdown(f'[바로가기]("https://arxiv.org/abs/2309.15217")', unsafe_allow_html=True)

current_dir = os.path.dirname(os.path.abspath(__file__))
image_path = os.path.join(current_dir, "images", "faithfulness.jpg")
image_path2 = os.path.join(current_dir, "images", "amnswerrelevancy.jpg")
image_path3 = os.path.join(current_dir, "images", "contextrecall.jpg")
image_path4 = os.path.join(current_dir, "images", "contextprecision.jpg")



st.image(image_path)
st.image(image_path2)
st.image(image_path3)
st.image(image_path4)




st.markdown("* * *")
st.subheader("|개별점수")
# Set the current directory and file path
current_dir = os.path.dirname(os.path.abspath(__file__))
file_path1 = os.path.join(current_dir, 'test_result.csv')

# Check if the file exists
if os.path.exists(file_path1):

    final_data = pd.read_csv(file_path1)
    # Display the DataFrame
    st.dataframe(final_data)


image_path5 = os.path.join(current_dir, "images", "test_final.jpg")
st.markdown("* * *")
st.subheader("|종합점수")
st.image(image_path5)


st.markdown("* * *")
st.subheader("|한계점")
st.write("1. llm(GPT4)으로 qna 셋을 만들어 그라운드 데이터가 정확하지 않다. ")
st.write("2. 알 수 없는 이유로 null값이 많이 출력되었다.(토큰 수 제한과 관련있어 보임)")
st.write("3. 유료모델(GPTAPI) 사용으로 테스트에 한계가 있었다.(테스트 한 번에 5-7달러 가량 소모)")
st.markdown("* * *")
st.subheader("|앞으로의 과제")
st.write("1. qna 데이터셋을 직접 만들어 오차율을 줄일 것. ")
st.write("2. null값 줄이기. ")
st.write("3. ollama모델로 대체할 방법 찾기. ")
