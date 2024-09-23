import streamlit as st
import PyPDF2
import docx  # For reading DOCX files
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain_openai import ChatOpenAI
from langchain.schema import Document
from string import Template
import re
import os
import time
from langchain.prompts import ChatPromptTemplate
from funcs import load_css, load_local_font

current_dir = os.path.dirname(os.path.abspath(__file__))
css_path = os.path.join(current_dir, 'style.css')

# CSS 파일 로드 함수
def load_css(file_name):
    if os.path.exists(file_name):
        with open(file_name, 'r', encoding='utf-8') as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    else:
        st.error(f"CSS 파일을 찾을 수 없습니다: {file_name}")

# CSS 파일 로드
load_css(css_path)

# Pretendard 폰트 로드
st.markdown("""
    <link href="https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css" rel="stylesheet">
    <style>
        html, body, [class*="css"] {
            font-family: 'Pretendard', sans-serif;
        }
    </style>
    """, unsafe_allow_html=True)

# API 키 입력 및 저장
st.sidebar.title("설정")
api_key = st.sidebar.text_input("OpenAI API Key 입력", type="password")

# API 키를 세션 상태에 저장
if api_key:
    st.session_state.api_key = api_key
    st.sidebar.success("API Key가 설정되었습니다.")
else:
    st.session_state.api_key = None

# PDF 파일 읽기 함수
def read_pdf(file_path):
    pdf_text = []
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            pdf_text.append(page.extract_text())
    return pdf_text

# DOCX 파일 읽기 함수
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# 파일 경로
docx_file_path = "./data/recipes.docx"
pdf_file_path2 = "./1.pdf"

# PDF 및 DOCX 파일에서 텍스트 불러오기
docx_content = read_docx(docx_file_path)
pdf_content2 = read_pdf(pdf_file_path2)

import pandas as pd

# 데이터를 딕셔너리 형태로 정리
data = {
    "항목": [
        "이름", "만나이", "성별", "키", "체중", "허리둘레", 
        "이상지질혈증 여부", "당뇨병 여부", 
        "음주 여부", "음주 빈도", "1회 주량", 
        "폭음 횟수", "절주 권고", 
        "음주 상담 유무",
        "고강도 운동 여부", "1주일", "1회 운동 시간", 
        "중강도 운동 여부", "1주일", "1회 운동 시간?", 
        "걷기, 자전거 운동", "1주일", "총 운동 시간",
        "음주 점수", "신체활동 점수", "고혈압 확률",
    ],
    "값": [
        "김첨지", 38, "남자", "173.0 cm", "89.0 kg", "88.9 cm",
        "없음", "없음",
        "마신다", "월1회정도", "1-2잔",
        "월 1회 미만", "없음", "없음",
        "안 한다", "None", "None", 
        "안 한다", "None", "None",
        "한다", "5일", "2시간 0분",
        "9점", "6.39859점", "21.41%",
    ]
}

# 데이터프레임 생성
df = pd.DataFrame(data)
df1 = df.reset_index(drop=True)

# 사이드바에 데이터프레임 표시
st.sidebar.title("김첨지님의 건강 보고서")
st.sidebar.dataframe(df1)

# PDF 및 DOCX 데이터를 문서 리스트로 변환
documents = []

# docx_content는 하나의 문자열이므로 바로 추가
if isinstance(docx_content, str):
    documents.append(Document(page_content=docx_content))
else:
    st.error("DOCX 내용이 올바른 문자열 형식이 아닙니다.")

# pdf_content2도 리스트일 가능성이 있으므로 개별 텍스트로 변환
for text in pdf_content2:
    if isinstance(text, str):
        documents.append(Document(page_content=text))
    else:
        st.error("두 번째 PDF 내용이 올바른 문자열 형식이 아닙니다.")

# 텍스트를 청크로 분리하는 함수
def get_text_chunks(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=512,  # 청크 크기를 더 줄임
        chunk_overlap=50,
        length_function=len
    )
    chunks = text_splitter.split_documents(documents)
    return chunks

# 벡터 스토어 생성 함수
def get_vectorstore(_text_chunks):
    embeddings = HuggingFaceEmbeddings(
        model_name="jhgan/ko-sroberta-multitask",
        model_kwargs={'device': 'cuda'},
        encode_kwargs={'normalize_embeddings': True}
    )
    vectordb = FAISS.from_documents(_text_chunks, embeddings)
    return vectordb

# 텍스트 청크 생성
text_chunks = get_text_chunks(documents)

# 벡터 스토어 생성
vectorstore = get_vectorstore(text_chunks)

# 대화형 체인 생성 함수
def get_conversation_chain(_vectorstore):
    retriever = _vectorstore.as_retriever(
        search_type='similarity',
        search_kwargs={"k": 10},
        verbose=False
    )

    # Check if API key is set
    if not st.session_state.get('api_key'):
        st.error("API Key를 입력해 주세요.")
        return None

    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=ChatOpenAI(openai_api_key=st.session_state.api_key), 
        chain_type="stuff",
        retriever=retriever, 
        return_source_documents=True,
        verbose=False
    )
    return conversation_chain

# 대화형 체인 생성
conversation_chain = get_conversation_chain(vectorstore)

if 'conversation' not in st.session_state and conversation_chain:
    st.session_state.conversation = conversation_chain

# 사용자 질의 템플릿 생성 함수
def get_prompt_template(query):
    # 템플릿 작성에 대한 로직 (예: 식단, 레시피 관련)
    prompt = Template(f"""
    {docx_file_path} 자료 또는 추가적인 정보를 포함해서 자유롭게 답변해 줘.
    {pdf_content2}의 내용을 토대로 대답해 줘.
    """)
    return prompt.substitute(query=query)

# Chat logic
if query := st.chat_input("식단 또는 레시피를 물어보세요.") and st.session_state.get('conversation'):
    formatted_query = get_prompt_template(query)

    with st.chat_message("user"):
        st.markdown(query)

    with st.chat_message("assistant"):
        chain = st.session_state.conversation

        with st.spinner("Thinking..."):
            result = chain({"question": formatted_query, "chat_history": []})
            response = result['answer']
            source_documents = result['source_documents']

            response_container = st.empty()
            chunk_size = 5
            for i in range(0, len(response), chunk_size):
                response_container.markdown(response[:i + chunk_size], unsafe_allow_html=True)
                time.sleep(0.05)

            with st.expander("참고 문서 확인"):
                for doc in source_documents:
                    st.markdown(doc.metadata.get('source', ''), help=doc.page_content)

# 캐시 지우기 버튼
if st.button('캐시 지우기'):
    st.cache_data.clear()
    st.cache_resource.clear()
    st.success("캐시가 지워졌습니다!")
